"""
Render CV dict → DOCX (python-docx, ATS-native format).

Optimized for Workday/SuccessFactors parsers: single column, no text boxes,
no fancy styling, standard section headers.
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .inline_markdown import iter_docx_runs


# Color constants
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
GRAY_LIGHT = RGBColor(0x88, 0x88, 0x88)


def _add_inline_runs(paragraph, text: str, size_pt: float = 10.5,
                     color: RGBColor = BLACK) -> None:
    """Add inline-markdown-aware runs to a paragraph (preserves bold/italic)."""
    for display, bold, italic in iter_docx_runs(text or ""):
        run = paragraph.add_run(display)
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def render_docx(cv: dict, output_path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Document-wide settings
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Normal style: Calibri 10.5pt
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK

    # ── HEADER ──
    _write_header(doc, cv.get("header", {}))

    # ── SECTIONS ──
    for sec in cv.get("sections", []):
        _write_section(doc, sec)

    # ── METADATA (core properties) ──
    cp = doc.core_properties
    cp.author = cv.get("header", {}).get("name", "")
    cp.title = f"{cv.get('header', {}).get('name', '')} — " \
               f"{cv.get('header', {}).get('headline', '')}"
    cp.subject = cv.get("header", {}).get("headline", "")
    cp.keywords = ", ".join(filter(None, [
        cv.get("meta", {}).get("target_role"),
        cv.get("meta", {}).get("target_company"),
        cv.get("meta", {}).get("recipe"),
    ]))

    doc.save(str(output_path))
    return output_path


def _write_header(doc, header: dict):
    # Name (big, bold, centered)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(header.get("name", ""))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BLACK

    # Headline (medium, centered)
    if header.get("headline"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header["headline"])
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY

    # Contact lines (centered)
    for line in header.get("contact_lines", []):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY

    # Horizontal separator (added as space)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(6)


def _write_section(doc, sec: dict):
    # Section title (uppercase, with bottom border)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(sec.get("title", "").upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    _add_bottom_border(p)

    stype = sec.get("type")
    if stype == "paragraph":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_inline_runs(p, sec.get("content", ""), size_pt=10.5)

    elif stype == "experience":
        for entry in sec.get("entries", []):
            _write_job_entry(doc, entry)

    elif stype == "skills":
        for cat in sec.get("categories", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            if cat.get("name"):
                run = p.add_run(f"{cat['name']}: ")
                run.bold = True
                run.font.size = Pt(10.5)
            # Items may contain inline markdown; render each then comma-join via
            # alternating runs so bold/italic survive.
            for idx, item in enumerate(cat.get("items", [])):
                if idx > 0:
                    sep = p.add_run(", ")
                    sep.font.size = Pt(10.5)
                _add_inline_runs(p, item, size_pt=10.5)

    elif stype == "list":
        for item in sec.get("items", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            _add_inline_runs(p, item, size_pt=10.5)


def _write_job_entry(doc, entry: dict):
    # Company — Role (bold)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(entry.get("company", ""))
    run.bold = True
    run.font.size = Pt(11)
    if entry.get("role"):
        run = p.add_run(" — ")
        run.font.color.rgb = GRAY_LIGHT
        run.font.size = Pt(11)
        run = p.add_run(entry["role"])
        run.bold = True
        run.font.size = Pt(11)

    # Dates · Location (italic gray)
    meta = entry.get("dates", "")
    if entry.get("location"):
        meta = f"{meta} · {entry['location']}"
    if meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(meta)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY

    # Bullets
    for b in entry.get("bullets", []):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _add_inline_runs(p, b, size_pt=10.5)


def _add_bottom_border(paragraph):
    """Add a thin gray bottom border to a paragraph (section title underline)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
