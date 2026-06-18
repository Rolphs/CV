"""
Audit historical CVs (2010-2022) and extract candidate new info for master.

Output: data/reports/historical_audit.md (human-readable analysis)
       data/reports/historical_extracts.json (raw text per file)

Focus:
  - Extract full text from each PDF/DOCX
  - Group by year (heuristic from filename)
  - Identify NEW content vs. master:
      * Pre-2007 puestos (early career)
      * Educación / formación académica
      * Certificaciones (Miller Heiman, etc.)
      * Habilidades / metodologías mencionadas
      * Logros con métricas no capturados
      * Clientes / proyectos específicos
"""
from __future__ import annotations
import json
import re
import unicodedata
from collections import defaultdict
from pathlib import Path
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parent.parent
HIST = ROOT / "data" / "raw" / "06_versiones_historicas_2010_2022"
MASTER = ROOT / "data" / "master" / "cv_master_raul_mercado.xlsx"
OUT = ROOT / "data" / "reports"


def normalize(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c)).lower()


def extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        r = PdfReader(str(path))
        return "\n".join(p.extract_text() or "" for p in r.pages)
    except Exception as e:
        return f"[ERROR PDF: {e}]"


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
        d = Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        # tables too
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as e:
        return f"[ERROR DOCX: {e}]"


def extract(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    return f"[SKIP: unsupported extension {ext}]"


# ─── Year detection ─────────────────────────────────────────────────────────
def detect_year(filename: str) -> str:
    m = re.search(r"\b(20\d{2})\b", filename)
    return m.group(1) if m else "?"


# ─── Master snapshot ────────────────────────────────────────────────────────
def load_master_snapshot():
    wb = load_workbook(str(MASTER), data_only=True)
    empresas = []
    for row in wb["02 Experiencia"].iter_rows(min_row=3, max_col=10, values_only=True):
        if row[0] is None:
            continue
        empresas.append({
            "id": int(row[0]), "name": row[1], "puesto": row[4],
            "fecha_ini": row[8], "fecha_fin": row[9],
        })
    logros_es = []
    for row in wb["03 Logros"].iter_rows(min_row=3, max_col=14, values_only=True):
        if row[0] is None:
            continue
        logros_es.append(normalize(str(row[4] or "")))
    wb.close()
    return empresas, logros_es


# ─── Pattern detection ──────────────────────────────────────────────────────
EDUCATION_PATTERNS = [
    r"(?i)\b(universidad|UNAM|UAM|ITAM|IBERO|TEC[\b\s])",
    r"(?i)\b(licenciatura|maestr[íi]a|doctorado|diplomado|carrera|estudios|bachelor|master|MBA)\b",
    r"(?i)\b(comunicaci[óo]n|psicolog[íi]a|mercadotecnia|administraci[óo]n|ingenier[íi]a|sociolog[íi]a|antropolog[íi]a|electr[óo]nica)\b",
]
CERT_PATTERNS = [
    r"(?i)\b(certificaci[óo]n|certificado|certified|curso|workshop|diplomado)\b",
    r"(?i)\b(Miller\s*Heiman|Strategic\s*Selling|Needscope|SPSS|Morae|ATLAS\.?ti|NVivo|tableau|power\s*bi)\b",
]
EARLY_JOB_PATTERNS = [
    r"(?i)\b(19[89]\d|200[0-6])\b",  # años 1980-2006
]
CLIENT_PATTERNS = [
    r"(?i)\b(Unilever|Nestl[ée]|P&G|Procter|Coca[-\s]?Cola|Pepsi|Walmart|Banamex|Banorte|BBVA|Santander|Telmex|Telcel|AT&T|IBM|Microsoft|Google|HSBC|HP|Cinepolis|Cinemex|Sanborns|Liverpool|Bimbo|Modelo|FEMSA|Maseca|Herdez|La\s*Costen[aâ])\b",
]
METRIC_RE = re.compile(r"\d+\s*(?:%|x|mm?|k\b|m\b|millones?|million|m\$|MXN|USD)|\$\s*[\d.,]+|[+\-]\s*\d+")


def find_patterns(text: str, patterns: list[str]) -> set[str]:
    hits = set()
    for pat in patterns:
        for m in re.finditer(pat, text):
            hits.add(m.group(0).strip())
    return hits


def main():
    print(f"📂 Scanning {HIST}")
    files = sorted(
        [f for f in HIST.rglob("*") if f.is_file() and f.suffix.lower() in (".pdf", ".docx")],
        key=lambda p: (detect_year(p.name), p.name),
    )
    print(f"Found {len(files)} extractable files (PDF/DOCX)")

    master_empresas, master_logros_norm = load_master_snapshot()
    master_company_names = {normalize(e["name"]) for e in master_empresas}
    print(f"Master has {len(master_empresas)} puestos, {len(master_logros_norm)} logros")

    extracts = {}
    for f in files:
        rel = f.relative_to(HIST).as_posix()
        text = extract(f)
        extracts[rel] = {
            "year": detect_year(f.name),
            "size_kb": round(f.stat().st_size / 1024, 1),
            "text_len": len(text),
            "text": text,
        }
        status = "✓" if len(text) > 200 else "⚠ poco texto"
        print(f"  {status}  [{detect_year(f.name)}]  {len(text):>6}ch  {rel}")

    # save raw
    raw_out = OUT / "historical_extracts.json"
    raw_out.write_text(json.dumps(
        {k: {kk: vv for kk, vv in v.items() if kk != "text" or len(v["text"]) < 50000}
         for k, v in extracts.items()},
        indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\n📁 Raw extracts (truncated) → {raw_out.name}")

    # Build analysis
    by_year = defaultdict(list)
    for rel, data in extracts.items():
        by_year[data["year"]].append((rel, data))

    md = OUT / "historical_audit.md"
    with open(md, "w", encoding="utf-8") as f:
        f.write("# 📜 Auditoría de CVs Históricos 2010-2022\n\n")
        f.write(f"**Archivos procesados:** {len(extracts)}\n\n")
        f.write(f"**Maestro actual:** {len(master_empresas)} puestos, {len(master_logros_norm)} logros\n\n")

        f.write("## 📊 Inventario por año\n\n")
        f.write("| Año | # archivos | Total chars extraídos |\n|---|---|---|\n")
        for year in sorted(by_year.keys()):
            total = sum(d["text_len"] for _, d in by_year[year])
            f.write(f"| {year} | {len(by_year[year])} | {total:,} |\n")

        # Per-CV deep analysis
        f.write("\n\n## 🔍 Análisis por archivo\n\n")
        all_education = set()
        all_certs = set()
        all_clients = set()
        all_early_years = set()
        all_metrics = set()
        all_companies_mentioned = defaultdict(int)

        for year in sorted(by_year.keys()):
            f.write(f"\n### Año {year}\n\n")
            for rel, data in by_year[year]:
                if data["text_len"] < 200:
                    f.write(f"#### `{rel}` ({data['size_kb']} KB) — _poco texto, posible visual_\n\n")
                    continue
                text = data["text"]
                f.write(f"#### `{rel}` ({data['size_kb']} KB, {data['text_len']:,} chars)\n\n")

                # Detect mentions
                edu = find_patterns(text, EDUCATION_PATTERNS)
                certs = find_patterns(text, CERT_PATTERNS)
                clients = find_patterns(text, CLIENT_PATTERNS)
                early_years = set(re.findall(r"\b(19[89]\d|200[0-6])\b", text))
                metrics_found = METRIC_RE.findall(text)

                all_education.update(edu)
                all_certs.update(certs)
                all_clients.update(clients)
                all_early_years.update(early_years)
                all_metrics.update(metrics_found)

                # Companies mentioned
                for emp in master_empresas:
                    if normalize(emp["name"]).split()[0] in normalize(text):
                        all_companies_mentioned[emp["name"]] += 1

                if edu:
                    f.write(f"- **Educación:** {', '.join(sorted(edu)[:8])}\n")
                if certs:
                    f.write(f"- **Certs/tools:** {', '.join(sorted(certs)[:8])}\n")
                if clients:
                    f.write(f"- **Clientes/marcas:** {', '.join(sorted(clients)[:10])}\n")
                if early_years:
                    f.write(f"- **Años pre-2007 mencionados:** {', '.join(sorted(early_years))}\n")
                if metrics_found:
                    f.write(f"- **Métricas:** {', '.join(metrics_found[:8])}\n")
                f.write("\n")

        # Summary
        f.write("\n\n## 🎯 RESUMEN DE HALLAZGOS NUEVOS POTENCIALES\n\n")
        f.write("### Educación / Formación detectada\n")
        for e in sorted(all_education):
            f.write(f"- {e}\n")
        f.write("\n### Certificaciones / Herramientas detectadas\n")
        for c in sorted(all_certs):
            f.write(f"- {c}\n")
        f.write("\n### Marcas/Clientes mencionados\n")
        for c in sorted(all_clients):
            f.write(f"- {c}\n")
        f.write("\n### Años pre-2007 mencionados (posibles puestos tempranos)\n")
        for y in sorted(all_early_years):
            f.write(f"- {y}\n")

    print(f"📁 Analysis → {md.name}")


if __name__ == "__main__":
    main()
