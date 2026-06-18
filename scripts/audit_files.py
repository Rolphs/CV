"""
Audit & catalog every file in cv_rmb/source/.

For each file:
  - basic metadata (path, size, mtime, type)
  - extracted text preview (first 800 chars)
  - structural counts (pages for PDF, paragraphs for DOCX, sheets/rows for XLSX)
  - SHA1 hash for dedupe analysis

Outputs:
  - data/reports/file_catalog.json   (machine-readable, full)
  - data/reports/file_catalog.md     (human-readable summary)
  - data/reports/duplicates.md       (groups of files with identical hash)
"""

from __future__ import annotations

import hashlib
import json
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "data" / "raw"
OUT = ROOT / "data" / "reports"
OUT.mkdir(exist_ok=True)

PREVIEW_CHARS = 800


# ─── extractors ──────────────────────────────────────────────────────────────

def extract_pdf(path: Path) -> dict:
    try:
        reader = PdfReader(str(path))
        pages = len(reader.pages)
        text_chunks = []
        for page in reader.pages:
            try:
                text_chunks.append(page.extract_text() or "")
            except Exception as exc:
                text_chunks.append(f"[page extract error: {exc}]")
        full_text = "\n".join(text_chunks).strip()
        return {
            "type": "pdf",
            "pages": pages,
            "char_count": len(full_text),
            "preview": full_text[:PREVIEW_CHARS],
            "full_text": full_text,
        }
    except Exception as exc:
        return {"type": "pdf", "error": str(exc)}


def extract_docx(path: Path) -> dict:
    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for tbl in doc.tables:
            for row in tbl.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    tables_text.append(" | ".join(row_cells))
        full_text = "\n".join(paragraphs + tables_text).strip()
        return {
            "type": "docx",
            "paragraphs": len(paragraphs),
            "table_rows": len(tables_text),
            "char_count": len(full_text),
            "preview": full_text[:PREVIEW_CHARS],
            "full_text": full_text,
        }
    except Exception as exc:
        return {"type": "docx", "error": str(exc)}


def extract_xlsx(path: Path) -> dict:
    try:
        wb = load_workbook(str(path), data_only=True, read_only=True)
        sheets_info = []
        all_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0
            sample_rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 5:
                    break
                sample_rows.append([str(c) if c is not None else "" for c in row])
            sheets_info.append({
                "name": sheet_name,
                "max_row": max_row,
                "max_col": max_col,
                "sample_first_5_rows": sample_rows,
            })
            all_text.append(f"[{sheet_name}] {max_row}r x {max_col}c")
        wb.close()
        preview = " | ".join(all_text)
        return {
            "type": "xlsx",
            "sheets": sheets_info,
            "sheet_count": len(sheets_info),
            "preview": preview[:PREVIEW_CHARS],
        }
    except Exception as exc:
        return {"type": "xlsx", "error": str(exc)}


def extract_unknown(path: Path) -> dict:
    return {"type": "other", "preview": ""}


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_unknown,   # legacy .doc not supported by python-docx
    ".xlsx": extract_xlsx,
    ".xlsm": extract_xlsx,
}


# ─── walk + audit ────────────────────────────────────────────────────────────

def sha1(path: Path) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def audit_file(path: Path) -> dict:
    stat = path.stat()
    rel = path.relative_to(ROOT).as_posix()
    record = {
        "rel_path": rel,
        "name": path.name,
        "size_bytes": stat.st_size,
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "ext": path.suffix.lower(),
    }
    if stat.st_size == 0:
        record["status"] = "empty"
        record["sha1"] = None
        record["extract"] = {"type": "empty"}
        return record

    record["sha1"] = sha1(path)
    extractor = EXTRACTORS.get(path.suffix.lower(), extract_unknown)
    record["extract"] = extractor(path)
    record["status"] = "ok" if "error" not in record["extract"] else "error"
    return record


def main():
    records = []
    for path in sorted(SOURCE.rglob("*")):
        if path.is_file():
            safe = str(path.relative_to(SOURCE)).encode("ascii", "replace").decode("ascii")
            print(f"  ... {safe}")
            records.append(audit_file(path))

    # Full JSON dump
    json_path = OUT / "file_catalog.json"
    json_path.write_text(json.dumps(records, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\n✅ JSON catalog → {json_path}")

    # Dedupe groups (by sha1)
    by_hash: dict[str, list[dict]] = defaultdict(list)
    for r in records:
        if r.get("sha1"):
            by_hash[r["sha1"]].append(r)
    dup_groups = {h: rs for h, rs in by_hash.items() if len(rs) > 1}

    # Markdown human summary
    md_path = OUT / "file_catalog.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# 📋 CV-RMB · Catálogo de archivos\n\n")
        f.write(f"_Generado: {datetime.now().isoformat(timespec='seconds')}_\n\n")
        total = len(records)
        empty = sum(1 for r in records if r["status"] == "empty")
        errors = sum(1 for r in records if r["status"] == "error")
        f.write(f"**Total:** {total} archivos · **Vacíos (0B):** {empty} · **Errores:** {errors} · **Grupos duplicados:** {len(dup_groups)}\n\n")

        # Group by parent folder
        by_folder: dict[str, list[dict]] = defaultdict(list)
        for r in records:
            parent = str(Path(r["rel_path"]).parent)
            by_folder[parent].append(r)

        for folder in sorted(by_folder):
            f.write(f"\n## 📁 `{folder}`\n\n")
            for r in by_folder[folder]:
                size_kb = r["size_bytes"] / 1024
                badge = ""
                if r["status"] == "empty":
                    badge = " · ⚠️ vacío"
                elif r["status"] == "error":
                    badge = " · ❌ error"
                f.write(f"### `{r['name']}` ({size_kb:.1f} KB){badge}\n")
                f.write(f"- **mod:** {r['modified']}\n")
                ext = r["extract"]
                if ext.get("type") == "pdf":
                    f.write(f"- **páginas:** {ext.get('pages')} · **chars:** {ext.get('char_count')}\n")
                elif ext.get("type") == "docx":
                    f.write(f"- **párrafos:** {ext.get('paragraphs')} · **filas tabla:** {ext.get('table_rows')} · **chars:** {ext.get('char_count')}\n")
                elif ext.get("type") == "xlsx":
                    f.write(f"- **hojas:** {ext.get('sheet_count')}\n")
                    for sh in ext.get("sheets", []):
                        f.write(f"  - `{sh['name']}`: {sh['max_row']}r × {sh['max_col']}c\n")
                if ext.get("error"):
                    f.write(f"- **error:** `{ext['error']}`\n")
                preview = (ext.get("preview") or "").strip()
                if preview:
                    snippet = preview[:300].replace("\n", " ⏎ ")
                    f.write(f"- **preview:** {snippet}…\n")
                f.write("\n")

    print(f"✅ Markdown catalog → {md_path}")

    # Duplicates report
    dup_path = OUT / "duplicates.md"
    with open(dup_path, "w", encoding="utf-8") as f:
        f.write("# 🧬 Duplicados exactos (mismo SHA1)\n\n")
        if not dup_groups:
            f.write("_No se encontraron duplicados exactos._\n")
        else:
            for h, rs in sorted(dup_groups.items(), key=lambda kv: -len(kv[1])):
                f.write(f"\n## Grupo (sha1 `{h[:10]}…`) — {len(rs)} archivos\n")
                for r in rs:
                    f.write(f"- `{r['rel_path']}` ({r['size_bytes']/1024:.1f} KB)\n")
    print(f"✅ Duplicates report → {dup_path}")


if __name__ == "__main__":
    main()
